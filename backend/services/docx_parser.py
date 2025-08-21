import re
from typing import List, Dict, Any, Optional
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocxParser:
    def __init__(self):
        self.questions = []
        self.current_question = None
        self.current_choices = []
        self.file_metadata = {
            'subject': None,
            'lecturer': None
        }
    
    def parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse DOCX file và trả về danh sách questions"""
        try:
            doc = Document(file_path)
            self.questions = []
            self.current_question = None
            self.current_choices = []
            self.file_metadata = {
                'subject': None,
                'lecturer': None
            }
            
            logger.info(f"Parsing DOCX file: {file_path}")
            logger.info(f"Total paragraphs: {len(doc.paragraphs)}")
            logger.info(f"Total tables: {len(doc.tables)}")
            
            # Parse từng paragraph
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue
                
                logger.info(f"Paragraph {i}: '{text}'")
                
                # Parse file metadata first
                if self._is_file_metadata(text):
                    logger.info(f"  -> File metadata: {text}")
                    self._parse_file_metadata(text)
                    continue
                
                # Parse question
                if self._is_question_start(text):
                    logger.info(f"  -> Question start: {text}")
                    self._save_current_question()
                    self._start_new_question(text)
                # Parse choice
                elif self._is_choice(text):
                    logger.info(f"  -> Choice: {text}")
                    self._add_choice(text)
                # Parse answer
                elif self._is_answer(text):
                    logger.info(f"  -> Answer: {text}")
                    self._set_answer(text)
                # Parse other metadata
                elif self._is_question_metadata(text):
                    logger.info(f"  -> Question metadata: {text}")
                    self._set_question_metadata(text)
                # Add to question text
                else:
                    logger.info(f"  -> Question text: {text}")
                    self._add_to_question_text(text)
            
            # Parse tables
            for i, table in enumerate(doc.tables):
                logger.info(f"Parsing table {i} with {len(table.rows)} rows")
                self._parse_table(table)
            
            # Save last question
            self._save_current_question()
            
            logger.info(f"Parsed {len(self.questions)} questions")
            logger.info(f"File metadata: {self.file_metadata}")
            return self.questions
            
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {e}")
            raise
    
    def _is_file_metadata(self, text: str) -> bool:
        """Kiểm tra có phải metadata của file không"""
        metadata_patterns = [
            r'^Subject:\s*',
            r'^Lecturer:\s*'
        ]
        return any(re.match(pattern, text) for pattern in metadata_patterns)
    
    def _parse_file_metadata(self, text: str):
        """Parse metadata của file"""
        # Parse Subject
        subject_match = re.match(r'^Subject:\s*(.+)', text)
        if subject_match:
            self.file_metadata['subject'] = subject_match.group(1).strip()
            logger.info(f"Found subject: {self.file_metadata['subject']}")
        
        # Parse Lecturer
        lecturer_match = re.match(r'^Lecturer:\s*(.+)', text)
        if lecturer_match:
            self.file_metadata['lecturer'] = lecturer_match.group(1).strip()
            logger.info(f"Found lecturer: {self.file_metadata['lecturer']}")
    
    def get_file_metadata(self) -> Dict[str, Any]:
        """Lấy metadata của file"""
        return self.file_metadata.copy()
    

    
    def _is_question_start(self, text: str) -> bool:
        """Kiểm tra có phải bắt đầu câu hỏi không"""
        return re.match(r'^QN=\d+', text) is not None
    
    def _is_choice(self, text: str) -> bool:
        """Kiểm tra có phải choice không"""
        return re.match(r'^[a-d]\.\s', text) is not None
    
    def _is_answer(self, text: str) -> bool:
        """Kiểm tra có phải answer không"""
        return re.match(r'^ANSWER:\s*[A-D]', text) is not None
    
    def _is_question_metadata(self, text: str) -> bool:
        """Kiểm tra có phải metadata của câu hỏi không"""
        metadata_patterns = [
            r'^MARK:\s*',
            r'^UNIT:\s*',
            r'^MIX CHOICES:\s*'
        ]
        return any(re.match(pattern, text) for pattern in metadata_patterns)
    
    def _start_new_question(self, text: str):
        """Bắt đầu câu hỏi mới"""
        # Extract question number
        match = re.match(r'QN=(\d+)', text)
        question_number = int(match.group(1)) if match else 1
        
        self.current_question = {
            'question_number': question_number,
            'question_text': '',
            'choices': [],
            'answer': None,
            'mark': 1.0,
            'unit': '',
            'mix_choices': True,
            'image': None
        }
        self.current_choices = []
    
    def _add_to_question_text(self, text: str):
        """Thêm text vào câu hỏi"""
        if self.current_question:
            # Check for image reference
            image_match = re.search(r'\[file:([^\]]+)\]', text)
            if image_match:
                self.current_question['image'] = image_match.group(1)
                # Remove image reference from text
                text = re.sub(r'\[file:[^\]]+\]', '', text).strip()
            
            if text:
                if self.current_question['question_text']:
                    self.current_question['question_text'] += ' ' + text
                else:
                    self.current_question['question_text'] = text
    
    def _add_choice(self, text: str):
        """Thêm choice"""
        if self.current_question:
            # Extract choice letter and content
            match = re.match(r'^([a-d])\.\s*(.+)', text)
            if match:
                choice_letter = match.group(1)
                choice_content = match.group(2).strip()
                
                choice = {
                    'letter': choice_letter,
                    'content': choice_content,
                    'is_correct': False
                }
                self.current_choices.append(choice)
    
    def _set_answer(self, text: str):
        """Set answer cho câu hỏi"""
        if self.current_question:
            match = re.match(r'^ANSWER:\s*([A-D])', text)
            if match:
                answer_letter = match.group(1).lower()
                self.current_question['answer'] = answer_letter
                
                # Mark correct choice
                for choice in self.current_choices:
                    if choice['letter'] == answer_letter:
                        choice['is_correct'] = True
                        break
    
    def _set_question_metadata(self, text: str):
        """Set metadata cho câu hỏi"""
        if not self.current_question:
            return
        
        # Parse MARK
        mark_match = re.match(r'^MARK:\s*([\d.]+)', text)
        if mark_match:
            self.current_question['mark'] = float(mark_match.group(1))
        
        # Parse UNIT
        unit_match = re.match(r'^UNIT:\s*(.+)', text)
        if unit_match:
            self.current_question['unit'] = unit_match.group(1).strip()
        
        # Parse MIX CHOICES
        mix_match = re.match(r'^MIX CHOICES:\s*(Yes|No)', text)
        if mix_match:
            self.current_question['mix_choices'] = mix_match.group(1).lower() == 'yes'
    
    def _save_current_question(self):
        """Lưu câu hỏi hiện tại"""
        if self.current_question and self.current_choices:
            # Validate question
            if not self.current_question['question_text']:
                logger.warning(f"Question {self.current_question['question_number']} has no text")
                return
            
            if len(self.current_choices) < 2:
                logger.warning(f"Question {self.current_question['question_number']} has less than 2 choices")
                return
            
            if not self.current_question['answer']:
                logger.warning(f"Question {self.current_question['question_number']} has no answer")
                return
            
            # Check for empty choices
            empty_choices = []
            for choice in self.current_choices:
                if not choice.get('content', '').strip():
                    empty_choices.append(choice['letter'].upper())
            
            if empty_choices:
                logger.warning(f"Question {self.current_question['question_number']} has empty choice(s): {', '.join(empty_choices)}")
                return
            
            # Convert to standard format
            question_data = {
                'question_number': self.current_question['question_number'],
                'question_text': self.current_question['question_text'],
                'choices': self.current_choices,
                'answer': self.current_question['answer'],
                'mark': self.current_question['mark'],
                'unit': self.current_question['unit'],
                'mix_choices': self.current_question['mix_choices'],
                'image': self.current_question['image']
            }
            
            self.questions.append(question_data)
    
    def _parse_table(self, table):
        """Parse table để tìm câu hỏi"""
        logger.info(f"Parsing table with {len(table.rows)} rows and {len(table.columns)} columns")
        
        current_question_data = None
        
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) < 2:
                continue
                
            cell0_text = row.cells[0].text.strip()
            cell1_text = row.cells[1].text.strip()
            
            logger.info(f"Table row {row_idx}: [{cell0_text}] | [{cell1_text}]")
            
            # Parse question start (QN=1)
            if cell0_text.startswith('QN='):
                # Save previous question if exists
                if current_question_data:
                    self._save_question_from_table(current_question_data)
                
                # Start new question
                question_num = int(cell0_text.split('=')[1])
                current_question_data = {
                    'question_number': question_num,
                    'question_text': cell1_text,
                    'choices': [],
                    'answer': None,
                    'mark': 1.0,
                    'unit': '',
                    'mix_choices': True,
                    'image': None
                }
                logger.info(f"  -> Started question {question_num}: {cell1_text}")
                
            # Parse choice (a., b., c., d.)
            elif cell0_text in ['a.', 'b.', 'c.', 'd.']:
                if current_question_data:
                    choice_letter = cell0_text[0]
                    choice = {
                        'letter': choice_letter,
                        'content': cell1_text,
                        'is_correct': False
                    }
                    current_question_data['choices'].append(choice)
                    logger.info(f"  -> Added choice {choice_letter}: {cell1_text}")
                    
            # Parse answer (ANSWER:)
            elif cell0_text == 'ANSWER:':
                if current_question_data:
                    answer_letter = cell1_text.upper()
                    current_question_data['answer'] = answer_letter.lower()
                    
                    # Mark correct choice
                    for choice in current_question_data['choices']:
                        if choice['letter'] == answer_letter.lower():
                            choice['is_correct'] = True
                            break
                    logger.info(f"  -> Set answer: {answer_letter}")
                    
            # Parse mark (MARK:)
            elif cell0_text == 'MARK:':
                if current_question_data:
                    try:
                        current_question_data['mark'] = float(cell1_text)
                        logger.info(f"  -> Set mark: {cell1_text}")
                    except ValueError:
                        logger.warning(f"Invalid mark value: {cell1_text}")
                        
            # Parse unit (UNIT:)
            elif cell0_text == 'UNIT:':
                if current_question_data:
                    current_question_data['unit'] = cell1_text
                    logger.info(f"  -> Set unit: {cell1_text}")
                    
            # Parse mix choices (MIX CHOICES:)
            elif cell0_text == 'MIX CHOICES:':
                if current_question_data:
                    current_question_data['mix_choices'] = cell1_text.lower() == 'yes'
                    logger.info(f"  -> Set mix choices: {cell1_text}")
        
        # Save last question
        if current_question_data:
            self._save_question_from_table(current_question_data)
    
    def _save_question_from_table(self, question_data):
        """Lưu câu hỏi từ table data"""
        # Validate question
        if not question_data['question_text']:
            logger.warning(f"Question {question_data['question_number']} has no text")
            return
        
        if len(question_data['choices']) < 2:
            logger.warning(f"Question {question_data['question_number']} has less than 2 choices")
            return
        
        if not question_data['answer']:
            logger.warning(f"Question {question_data['question_number']} has no answer")
            return
        
        # Check for empty choices
        empty_choices = []
        for choice in question_data['choices']:
            if not choice.get('content', '').strip():
                empty_choices.append(choice['letter'].upper())
        
        if empty_choices:
            logger.warning(f"Question {question_data['question_number']} has empty choice(s): {', '.join(empty_choices)}")
            return
        
        # Check for image reference in question text
        image_match = re.search(r'\[file:([^\]]+)\]', question_data['question_text'])
        if image_match:
            question_data['image'] = image_match.group(1)
            # Remove image reference from text
            question_data['question_text'] = re.sub(r'\[file:[^\]]+\]', '', question_data['question_text']).strip()
        
        # Convert to standard format
        question = {
            'question_number': question_data['question_number'],
            'question_text': question_data['question_text'],
            'choices': question_data['choices'],
            'answer': question_data['answer'],
            'mark': question_data['mark'],
            'unit': question_data['unit'],
            'mix_choices': question_data['mix_choices'],
            'image': question_data['image']
        }
        
        self.questions.append(question)
        logger.info(f"Saved question {question_data['question_number']} from table")
    
    def validate_questions(self) -> Dict[str, Any]:
        """Validate tất cả questions"""
        critical_errors = []  # Lỗi nghiêm trọng - không cho preview
        errors = []           # Lỗi thường - vẫn cho preview
        warnings = []
        
        # Kiểm tra metadata file
        if not self.file_metadata['subject']:
            critical_errors.append("Missing Subject information in file")
        
        # Kiểm tra có câu hỏi nào không
        if len(self.questions) == 0:
            critical_errors.append("No questions found in file")
            return {
                'valid': False,
                'critical_errors': critical_errors,
                'errors': errors,
                'warnings': warnings,
                'total_questions': 0
            }
        
        for question in self.questions:
            # Lỗi nghiêm trọng - không cho preview
            if not question['question_text']:
                critical_errors.append(f"Question {question['question_number']}: Missing question text")
            
            if len(question['choices']) < 2:
                critical_errors.append(f"Question {question['question_number']}: Need at least 2 choices")
            
            if not question['answer']:
                critical_errors.append(f"Question {question['question_number']}: Missing answer")
            
            # Check answer exists in choices
            answer_exists = any(choice['letter'] == question['answer'] for choice in question['choices'])
            if not answer_exists:
                critical_errors.append(f"Question {question['question_number']}: Answer '{question['answer']}' not found in choices")
            
            # Check for duplicate choice letters
            choice_letters = [choice['letter'] for choice in question['choices']]
            if len(choice_letters) != len(set(choice_letters)):
                critical_errors.append(f"Question {question['question_number']}: Duplicate choice letters")
            
            # Check for empty choices
            empty_choices = []
            for i, choice in enumerate(question['choices']):
                if not choice.get('content', '').strip():
                    empty_choices.append(choice['letter'].upper())
            
            if empty_choices:
                if len(empty_choices) == 1:
                    critical_errors.append(f"Question {question['question_number']}: Empty choice {empty_choices[0]}")
                else:
                    critical_errors.append(f"Question {question['question_number']}: Empty choices {', '.join(empty_choices)}")
            
            # Lỗi thường - vẫn cho preview
            if question['mark'] <= 0:
                errors.append(f"Question {question['question_number']}: Invalid mark (must be > 0)")
            
            if not question['unit']:
                warnings.append(f"Question {question['question_number']}: No unit specified")
        
        return {
            'valid': len(critical_errors) == 0,
            'critical_errors': critical_errors,
            'errors': errors,
            'warnings': warnings,
            'total_questions': len(self.questions)
        } 